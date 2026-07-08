"""
文档加载器

支持读取多种文档格式为纯文本:
  - PDF  → pypdf
  - DOCX → python-docx
  - TXT/MD → 直接读取
"""

import os
from typing import Optional
from dataclasses import dataclass

from common.log import logger


@dataclass
class Document:
    """加载后的文档"""
    content: str
    filename: str
    filetype: str
    metadata: dict = None


class DocumentLoader:
    """文档加载器，根据扩展名自动选择解析方式"""

    @staticmethod
    def load(file_path: str) -> Optional[Document]:
        """
        加载文档文件，返回提取的文本内容

        :param file_path: 文件路径
        :return: Document 对象，失败返回 None
        """
        if not os.path.exists(file_path):
            logger.error(f"[DocumentLoader] 文件不存在: {file_path}")
            return None

        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()

        logger.info(f"[DocumentLoader] 加载文件: {filename}")

        try:
            if ext == ".pdf":
                return DocumentLoader._load_pdf(file_path, filename)
            elif ext == ".docx":
                return DocumentLoader._load_docx(file_path, filename)
            elif ext in [".txt", ".md", ".json", ".csv", ".yaml", ".yml"]:
                return DocumentLoader._load_text(file_path, filename, ext)
            else:
                # 未知格式，尝试按文本读取
                logger.warning(f"[DocumentLoader] 未知格式 {ext}，尝试按文本读取")
                return DocumentLoader._load_text(file_path, filename, ext)
        except Exception as e:
            logger.exception(f"[DocumentLoader] 读取文件失败: {filename}, {e}")
            return None

    @staticmethod
    def _load_pdf(file_path: str, filename: str) -> Optional[Document]:
        """读取 PDF 文件"""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"--- 第 {i+1} 页 ---\n{text}")

        content = "\n\n".join(pages)
        logger.info(f"[DocumentLoader] PDF 读取完成: {len(pages)} 页, {len(content)} 字符")

        return Document(
            content=content,
            filename=filename,
            filetype="pdf",
            metadata={"pages": len(pages)},
        )

    @staticmethod
    def _load_docx(file_path: str, filename: str) -> Optional[Document]:
        """读取 Word 文件"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        content = "\n".join(paragraphs)
        logger.info(f"[DocumentLoader] DOCX 读取完成: {len(paragraphs)} 段, {len(content)} 字符")

        return Document(
            content=content,
            filename=filename,
            filetype="docx",
            metadata={"paragraphs": len(paragraphs)},
        )

    @staticmethod
    def _load_text(file_path: str, filename: str, ext: str) -> Optional[Document]:
        """读取纯文本文件"""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        logger.info(f"[DocumentLoader] 文本读取完成: {len(content)} 字符")
        return Document(
            content=content,
            filename=filename,
            filetype=ext.lstrip("."),
        )